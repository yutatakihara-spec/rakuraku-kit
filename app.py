import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- 1. AIの設定（セキュリティ対策版） ---
# Streamlitの「secrets」という機能を使って、安全にAPIキーを読み込みます
if "GOOGLE_API_KEY" in st.secrets:
    # ネット公開時（Streamlit Cloud）はこちらが動く
    api_key = st.secrets["GOOGLE_API_KEY"]
else:
    # ローカルPCでテストする時用（後ほど設定するファイルから読み込む）
    api_key = "あなたのAPIキーをここに" 

genai.configure(api_key=api_key.strip(), transport='rest')

# --- 2. ドキュメント作成関数（PDFの代わり） ---
def create_docx(content, shop_name):
    doc = Document()
    
    # タイトルを追加
    doc.add_heading(f"{shop_name} 事業計画書", 0)
    
    # AIが生成した文章を1行ずつ追加
    for line in content.split('\n'):
        if line.startswith('###'): # 見出しっぽければ強調
            doc.add_heading(line.replace('###', '').strip(), level=2)
        elif line.startswith('##'):
            doc.add_heading(line.replace('##', '').strip(), level=1)
        else:
            doc.add_paragraph(line)
            
    # メモリ上にファイルを保存
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. 画面構成（事業計画書作成楽楽キット） ---
st.set_page_config(page_title="事業計画書作成楽楽キット", layout="wide")
st.set_page_config(page_title="事業計画書作成楽楽キット", page_icon="📝")

# サイドバーに簡単な使い方ガイドを追加
st.sidebar.markdown("""
### 🛠️ 使い方ガイド
1. **Step 1**: ビジネスの基本を入力
2. **Step 2**: お金のシミュレーション
3. **Step 3**: AIが計画書を執筆
4. **保存**: Googleドキュメント形式でDL
---
*Produced by 楽楽キット チーム*
""")
st.title("🚀 事業計画書作成楽楽キット")

step = st.sidebar.radio("メニュー", ["Step 1: 基本情報", "Step 2: 収支分析", "Step 3: AI作成・保存"])

# セッション状態の初期化
if 'plan_text' not in st.session_state: st.session_state.plan_text = ""

# --- Step 1: 基本情報 ---
if step == "Step 1: 基本情報":
    st.header("📋 Step 1: 基本情報の入力")
    col1, col2 = st.columns(2)
    
    with col1:
        # 💡 選択肢に「その他」を追加
        st.session_state.industry = st.selectbox(
            "業種",["小売業", "飲食業", "サービス業", "建設業", "IT", "その他"], 
            index=0
        )
        
        # 💡 「その他」が選ばれた時だけ、詳細を書く欄を出す
        if st.session_state.industry == "その他":
            st.session_state.other_industry = st.text_input(
                "具体的な事業内容を教えてください", 
                value=st.session_state.get('other_industry', ""),
                placeholder="例：出張型のペット介護サービス"
            )
        else:
            st.session_state.other_industry = "" # その他以外なら空にしておく
            
        st.session_state.shop_name = st.text_input("屋号・会社名", value=st.session_state.get('shop_name', ""))
        
    with col2:
        st.session_state.target = st.text_area("ターゲット顧客", value=st.session_state.get('target', ""))
        st.session_state.strength = st.text_area("独自の強み", value=st.session_state.get('strength', ""))

# --- Step 2: 収支分析 ---
elif step == "Step 2: 収支分析":
    st.header("💰 Step 2: 収支シミュレーション")
    status = st.radio("現在の状況", ["起業の準備中", "既に事業を行っている"])
    col1, col2 = st.columns(2)
    with col1:
        money_val = st.number_input("初期投資額 または 借入額（万円）", value=500)
        unit_price = st.number_input("想定客単価（円）", value=2000)
    with col2:
        fixed_costs = st.number_input("月額固定費（家賃・人件費など計：万円）", value=40)
        variable_rate = st.slider("原価率（％）", 0, 100, 30)

    # 損益分岐点計算
    marginal_rate = 1 - (variable_rate / 100)
    breakeven = fixed_costs / marginal_rate if marginal_rate > 0 else 0
    st.session_state.finance_data = {"status": status, "money": money_val, "unit": unit_price, "breakeven": breakeven}
    
    st.info(f"あなたのビジネスの損益分岐点は、月間売上 **{breakeven:.1f} 万円** です。")

# --- Step 3: AI作成・保存 ---
elif step == "Step 3: AI作成・保存":
    st.header("📄 Step 3: 事業計画書の完成")
    
    if st.button("AIで事業計画書をフル生成"):
            with st.spinner("AIが清書しています..."):
                try:
                    model = genai.GenerativeModel('gemini-flash-latest')
                    f = st.session_state.finance_data
                    
                    # 💡 「その他」の場合は、入力された具体的な事業内容をAIに伝える
                    industry_name = st.session_state.industry
                    if industry_name == "その他" and st.session_state.other_industry:
                        industry_name = f"「{st.session_state.other_industry}」という事業"
                    
                    prompt = f"""
                    あなたは{industry_name}の専門コンサルタントです。
                    店名「{st.session_state.shop_name}」の事業計画書を作成してください。
                    
                    損益分岐点（月{f['breakeven']:.1f}万円）をどう超えるかの具体的な集客・営業戦略を詳しく書き、
                    最後は「以上」で終わらせてください。
                    """
                    response = model.generate_content(prompt)
                    st.session_state.plan_text = response.text
                except Exception as e:
                    st.error(f"エラー: {e}")

    if st.session_state.plan_text:
        st.markdown(st.session_state.plan_text)
        
        st.divider()
        st.subheader("📥 成果物のダウンロード")
        
        # Word（Googleドキュメント対応）ファイルの作成
        docx_data = create_docx(st.session_state.plan_text, st.session_state.shop_name)
        
        st.download_button(
            label="📄 Googleドキュメント形式で保存",
            data=docx_data,
            file_name=f"{st.session_state.shop_name}_事業計画書.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.info("※ダウンロードしたファイルを Googleドライブにアップロードして開けば、Googleドキュメントとして編集できます！")
